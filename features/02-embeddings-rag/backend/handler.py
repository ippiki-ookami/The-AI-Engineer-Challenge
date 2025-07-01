"""
Embeddings and RAG Handler - Complete implementation with document processing
"""
import sys
import asyncio
import os
import time
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import nest_asyncio
nest_asyncio.apply()

# Add base directory to path
sys.path.append(str(Path(__file__).parent.parent.parent.parent))

# Add aimakerspace directory to path (local to this backend)
current_dir = Path(__file__).parent
sys.path.append(str(current_dir))

from base.backend.base_handler import BaseChatHandler, BaseChatRequest
from base.backend.debug_logger import debug_track

# Import aimakerspace modules (now available in path)
from aimakerspace.text_utils import CharacterTextSplitter
from aimakerspace.vectordatabase import VectorDatabase
from aimakerspace.openai_utils.embedding import EmbeddingModel
from aimakerspace.openai_utils.chatmodel import ChatOpenAI
from aimakerspace.openai_utils.prompts import SystemRolePrompt, UserRolePrompt


class DocumentProcessor:
    """Handle different document types and extract text content"""
    
    @staticmethod
    def extract_text_from_file(file_content: bytes, file_name: str, content_type: str) -> str:
        """Extract text from different file types"""
        
        if content_type == 'text/plain':
            # Handle text files
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('latin-1')
                except UnicodeDecodeError:
                    raise ValueError("Could not decode text file")
        
        elif content_type == 'application/pdf':
            # Handle PDF files
            try:
                import PyPDF2
                import io
                
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text_parts.append(f"--- Page {page_num} ---\n{page_text}\n")
                    except Exception as e:
                        text_parts.append(f"--- Page {page_num} ---\n[Error extracting text from page: {str(e)}]\n")
                
                if not text_parts:
                    return f"[PDF File: {file_name}]\nNo extractable text found in PDF"
                
                full_text = "\n".join(text_parts).strip()
                return full_text
                
            except Exception as e:
                raise ValueError(f"Error processing PDF {file_name}: {str(e)}")
        
        elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            # Handle Word documents
            try:
                import docx
                import io
                
                doc = docx.Document(io.BytesIO(file_content))
                text_parts = []
                
                # Extract text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():  # Only add non-empty paragraphs
                        text_parts.append(para.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                if not text_parts:
                    return f"[Word Document: {file_name}]\nNo extractable text found in document"
                
                full_text = "\n".join(text_parts).strip()
                return full_text
                
            except Exception as e:
                raise ValueError(f"Error processing Word document {file_name}: {str(e)}")
        
        elif content_type == 'text/markdown':
            # Handle Markdown files
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('latin-1')
                except UnicodeDecodeError:
                    raise ValueError("Could not decode markdown file")
        
        else:
            raise ValueError(f"Unsupported file type: {content_type}. Supported types: TXT, PDF, DOC, DOCX, MD")


class RAGHandler(BaseChatHandler):
    """RAG (Retrieval-Augmented Generation) Handler with complete document processing"""
    
    def __init__(self):
        super().__init__("Embeddings and RAG")
        
        # Document storage
        self.uploaded_documents = []  # Store document metadata and content
        
        # RAG components (initialized when needed)
        self.text_splitter = None
        self.embedding_model = None
        self.vector_db = None
        self.chat_model = None
        
        # Chunk metadata tracking
        self.chunk_metadata = {}  # Maps chunk index to metadata
        self.chunk_to_doc_mapping = {}  # Maps chunk global index to doc info
        
        # Configuration
        self.config = {
            'chunk_size': 1000,
            'chunk_overlap': 200,
            'embedding_model': 'text-embedding-3-small',
            'chat_model': 'gpt-4o-mini',
            'temperature': 0.1,
            'search_k': 5,
            'rag_k': 4,
            'adjacent_chunks': 1,  # Number of chunks to include on each side
            'scoring_strategy': 'reembed',  # 'reembed', 'average', 'weighted', 'original'
            'comparison_chunk_size': 3000  # Larger chunk size for comparison
        }
        
        # Store alternative vector databases for comparison
        self.large_chunk_vector_db = None
        self.large_chunk_metadata = {}
    
    def process_file_content(self, file_content: bytes, file_name: str, content_type: str) -> str:
        """Process file content using DocumentProcessor"""
        return DocumentProcessor.extract_text_from_file(file_content, file_name, content_type)
    
    @debug_track("Processing Document Upload")
    async def process_document_upload(self, file_name: str, file_content: str, file_type: str, api_key: str) -> dict:
        """Process uploaded document and add to vector database"""
        
        # Create document entry
        document_id = f"{len(self.uploaded_documents)}_{int(time.time())}"
        document = {
            'id': document_id,
            'name': file_name,
            'type': file_type,
            'content': file_content,  # Store full content
            'size': len(file_content),
            'upload_time': time.time(),
            'preview': file_content[:500] + "..." if len(file_content) > 500 else file_content,
            'word_count': len(file_content.split()),
            'line_count': len(file_content.split('\n'))
        }
        
        # Add to document store
        self.uploaded_documents.append(document)
        
        # Initialize components if needed
        await self._initialize_components(api_key)
        
        # Process document into chunks
        chunks = await self._process_document_into_chunks(file_content)
        
        # Update vector database
        await self._update_vector_database()
        
        return {
            'success': True,
            'document_id': document_id,
            'document_name': file_name,
            'content_length': len(file_content),
            'chunks_created': len(chunks),
            'total_documents': len(self.uploaded_documents),
            'vector_db_ready': self.vector_db is not None and len(self.vector_db.vectors) > 0,
            'content_preview': document['preview']
        }
    
    @debug_track("Initializing RAG Components")
    async def _initialize_components(self, api_key: str):
        """Initialize all RAG components with API key"""
        # Set API key in environment
        os.environ["OPENAI_API_KEY"] = api_key
        
        # Initialize components
        if not self.text_splitter:
            self.text_splitter = CharacterTextSplitter(
                chunk_size=self.config['chunk_size'],
                chunk_overlap=self.config['chunk_overlap']
            )
        
        if not self.embedding_model:
            self.embedding_model = EmbeddingModel(self.config['embedding_model'])
        
        if not self.vector_db:
            self.vector_db = VectorDatabase(self.embedding_model)
        
        if not self.chat_model:
            self.chat_model = ChatOpenAI(self.config['chat_model'])
    
    @debug_track("Processing Document into Chunks")
    async def _process_document_into_chunks(self, content: str) -> List[str]:
        """Split document content into chunks"""
        if not self.text_splitter:
            raise ValueError("Text splitter not initialized")
        
        # Split the document into chunks
        chunks = self.text_splitter.split_texts([content])
        return chunks
    
    @debug_track("Updating Vector Database")
    async def _update_vector_database(self):
        """Rebuild vector database with all documents"""
        if not self.vector_db or not self.text_splitter:
            raise ValueError("Components not initialized")
        
        if not self.uploaded_documents:
            return
        
        # Clear previous metadata
        self.chunk_metadata = {}
        self.chunk_to_doc_mapping = {}
        
        all_chunks = []
        global_chunk_idx = 0
        
        # Process each document and track chunk positions
        for doc_idx, doc in enumerate(self.uploaded_documents):
            # Split document into chunks
            doc_chunks = self.text_splitter.split_texts([doc['content']])
            
            # Track metadata for each chunk
            for local_chunk_idx, chunk_text in enumerate(doc_chunks):
                # Store chunk metadata
                self.chunk_metadata[global_chunk_idx] = {
                    'doc_id': doc['id'],
                    'doc_name': doc['name'],
                    'doc_idx': doc_idx,
                    'local_chunk_idx': local_chunk_idx,
                    'total_chunks_in_doc': len(doc_chunks),
                    'global_chunk_idx': global_chunk_idx,
                    'chunk_text': chunk_text
                }
                
                # Store reverse mapping
                self.chunk_to_doc_mapping[global_chunk_idx] = {
                    'doc_id': doc['id'],
                    'local_idx': local_chunk_idx
                }
                
                all_chunks.append(chunk_text)
                global_chunk_idx += 1
        
        # Rebuild vector database with all chunks
        self.vector_db = await self.vector_db.abuild_from_list(all_chunks)
        
        # Also build large chunk vector database for comparison
        await self._build_large_chunk_vector_db()
    
    @debug_track("Building Large Chunk Vector Database")
    async def _build_large_chunk_vector_db(self):
        """Build a separate vector database with larger chunks for comparison"""
        if not self.embedding_model or not self.uploaded_documents:
            return
        
        # Create text splitter with larger chunk size
        large_splitter = CharacterTextSplitter(
            chunk_size=self.config['comparison_chunk_size'],
            chunk_overlap=self.config['chunk_overlap']  # Keep same overlap ratio
        )
        
        # Clear previous large chunk metadata
        self.large_chunk_metadata = {}
        
        all_large_chunks = []
        global_chunk_idx = 0
        
        # Process each document with large chunks
        for doc_idx, doc in enumerate(self.uploaded_documents):
            # Split document into large chunks
            large_chunks = large_splitter.split_texts([doc['content']])
            
            # Track metadata for each large chunk
            for local_chunk_idx, chunk_text in enumerate(large_chunks):
                self.large_chunk_metadata[global_chunk_idx] = {
                    'doc_id': doc['id'],
                    'doc_name': doc['name'],
                    'doc_idx': doc_idx,
                    'local_chunk_idx': local_chunk_idx,
                    'total_chunks_in_doc': len(large_chunks),
                    'global_chunk_idx': global_chunk_idx,
                    'chunk_text': chunk_text,
                    'chunk_size': len(chunk_text)
                }
                
                all_large_chunks.append(chunk_text)
                global_chunk_idx += 1
        
        # Build the large chunk vector database
        if all_large_chunks:
            from aimakerspace.vectordatabase import VectorDatabase
            self.large_chunk_vector_db = VectorDatabase(self.embedding_model)
            self.large_chunk_vector_db = await self.large_chunk_vector_db.abuild_from_list(all_large_chunks)
    
    @debug_track("Searching Documents")
    async def search_documents(self, query: str, api_key: str, k: int = None) -> list:
        """Search documents using vector database with optional adjacent chunk expansion"""
        if not self.vector_db or len(self.vector_db.vectors) == 0:
            return []
        
        k = k or self.config['search_k']
        adjacent = self.config.get('adjacent_chunks', 0)
        
        # Initial search
        initial_results = self.vector_db.search_by_text(query, k=k)
        
        if adjacent == 0:
            # No adjacent chunks requested, return normal results
            results = []
            for i, (text, score) in enumerate(initial_results, 1):
                # Find chunk metadata
                chunk_idx = self._find_chunk_index_by_text(text)
                metadata = self.chunk_metadata.get(chunk_idx, {})
                
                results.append({
                    'rank': i,
                    'text': text,
                    'similarity_score': float(score),
                    'full_text': text,
                    'chunk_info': {
                        'global_idx': chunk_idx,
                        'doc_name': metadata.get('doc_name', 'Unknown'),
                        'local_idx': metadata.get('local_chunk_idx', -1),
                        'total_in_doc': metadata.get('total_chunks_in_doc', -1)
                    }
                })
            return results
        
        # Process with adjacent chunks
        return await self._search_with_adjacent_chunks(query, initial_results, adjacent)
    
    def _find_chunk_index_by_text(self, text: str) -> int:
        """Find chunk index by matching text content"""
        for idx, metadata in self.chunk_metadata.items():
            if metadata['chunk_text'] == text:
                return idx
        return -1
    
    @debug_track("Searching with Adjacent Chunks")
    async def _search_with_adjacent_chunks(self, query: str, initial_results: list, adjacent: int) -> list:
        """Re-embed and re-score with adjacent chunks included"""
        expanded_results = []
        
        for i, (chunk_text, original_score) in enumerate(initial_results):
            # Find the chunk index
            chunk_idx = self._find_chunk_index_by_text(chunk_text)
            if chunk_idx == -1:
                continue
                
            metadata = self.chunk_metadata[chunk_idx]
            
            # Get adjacent chunks within the same document
            adjacent_chunks = self._get_adjacent_chunks(chunk_idx, adjacent)
            
            # Combine chunks into one text block
            combined_text = "\n\n".join([
                self.chunk_metadata[idx]['chunk_text'] 
                for idx in sorted(adjacent_chunks.keys())
            ])
            
            # Re-embed the combined text
            if self.config['scoring_strategy'] == 'reembed' and self.embedding_model:
                # Create new embedding for combined text
                combined_embedding = await self.embedding_model.async_get_embeddings([combined_text])
                query_embedding = await self.embedding_model.async_get_embeddings([query])
                
                # Calculate new similarity score
                import numpy as np
                new_score = float(np.dot(combined_embedding[0], query_embedding[0]))
            else:
                # Use original score for now (other strategies can be added)
                new_score = original_score
            
            expanded_results.append({
                'original_rank': i + 1,  # Keep track of original rank
                'rank': i + 1,  # Will be updated after re-sorting
                'text': chunk_text,
                'similarity_score': float(new_score),
                'original_score': float(original_score),
                'score_change': float(new_score - original_score),
                'score_change_pct': ((new_score - original_score) / original_score * 100) if original_score > 0 else 0,
                'full_text': combined_text,
                'chunk_info': {
                    'center_chunk': {
                        'global_idx': chunk_idx,
                        'doc_name': metadata.get('doc_name', 'Unknown'),
                        'local_idx': metadata.get('local_chunk_idx', -1),
                        'total_in_doc': metadata.get('total_chunks_in_doc', -1)
                    },
                    'included_chunks': list(adjacent_chunks.keys()),
                    'total_chunks_used': len(adjacent_chunks),
                    'adjacent_setting': adjacent
                }
            })
        
        # Re-sort by new scores if we re-embedded
        if self.config['scoring_strategy'] == 'reembed':
            expanded_results.sort(key=lambda x: x['similarity_score'], reverse=True)
            # Update ranks and calculate rank changes
            for i, result in enumerate(expanded_results):
                result['rank'] = i + 1
                result['rank_change'] = result['original_rank'] - result['rank']  # Positive means moved up
        
        return expanded_results
    
    def _get_adjacent_chunks(self, center_idx: int, adjacent: int) -> dict:
        """Get adjacent chunks from the same document"""
        center_metadata = self.chunk_metadata[center_idx]
        doc_id = center_metadata['doc_id']
        local_idx = center_metadata['local_chunk_idx']
        
        adjacent_chunks = {center_idx: center_metadata}
        
        # Find chunks from the same document
        for idx, metadata in self.chunk_metadata.items():
            if metadata['doc_id'] == doc_id:
                local_distance = abs(metadata['local_chunk_idx'] - local_idx)
                if 0 < local_distance <= adjacent:
                    adjacent_chunks[idx] = metadata
        
        return adjacent_chunks
    
    @debug_track("Getting Full Document Content")
    async def get_full_document_content(self, document_id: str) -> dict:
        """Get full document content with metadata for viewing"""
        # Find the document
        document = None
        for doc in self.uploaded_documents:
            if doc['id'] == document_id:
                document = doc
                break
        
        if not document:
            raise ValueError(f"Document with ID {document_id} not found")
        
        # Add start and end markers
        full_content = f"📄 === DOCUMENT START: {document['name']} ===\n\n"
        full_content += document['content']
        full_content += f"\n\n📄 === DOCUMENT END: {document['name']} ==="
        
        return {
            'document_id': document_id,
            'name': document['name'],
            'type': document['type'],
            'full_content': full_content,
            'original_content': document['content'],
            'size': document['size'],
            'word_count': document.get('word_count', len(document['content'].split())),
            'line_count': document.get('line_count', len(document['content'].split('\n'))),
            'upload_time': document['upload_time']
        }
    
    async def get_system_prompt(self) -> str:
        """Return system prompt for RAG"""
        return """You are a helpful assistant that answers questions based on provided context.
        
Use the context provided to answer questions accurately. If the context doesn't contain 
relevant information, say "I don't have enough information to answer that question."
        
Be concise and cite specific parts of the context when possible."""
    
    async def process_user_message(self, request: BaseChatRequest) -> str:
        """Process user message with RAG if documents are available"""
        user_message = request.user_message
        
        # If no documents, return normal message
        if not self.uploaded_documents:
            return f"{user_message}\n\n[Note: No documents uploaded yet. Upload documents to enable RAG functionality.]"
        
        # If vector database not ready, return message with note
        if not self.vector_db or len(self.vector_db.vectors) == 0:
            doc_names = [doc['name'] for doc in self.uploaded_documents]
            return f"{user_message}\n\n[Documents available: {', '.join(doc_names)}]\n[Note: Documents are being processed. Please try again in a moment.]"
        
        # Use RAG to answer the question
        try:
            # Search for relevant context
            search_results = self.vector_db.search_by_text(user_message, k=self.config['rag_k'])
            
            if not search_results:
                return f"{user_message}\n\n[Note: No relevant context found in uploaded documents.]"
            
            # Build context from search results
            context_parts = []
            for i, (text, score) in enumerate(search_results, 1):
                context_parts.append(f"[Context {i}]: {text}")
            
            context = "\n\n".join(context_parts)
            
            # Create enhanced message with context
            enhanced_message = f"""Context from uploaded documents:
{context}

Question: {user_message}

Please answer based on the provided context."""
            
            return enhanced_message
            
        except Exception as e:
            return f"{user_message}\n\n[Error accessing documents: {str(e)}]"
    
    async def enhance_messages(self, messages: list, request: BaseChatRequest) -> list:
        """Enhance messages with RAG context"""
        # For now, just return messages as-is since we handle RAG in process_user_message
        return messages
    
    @debug_track("Updating RAG Configuration")
    async def update_configuration(self, config_updates: dict) -> dict:
        """Update RAG configuration settings"""
        old_config = self.config.copy()
        
        # Update configuration
        for key, value in config_updates.items():
            if key in self.config:
                self.config[key] = value
        
        # Reinitialize components if needed
        if 'chunk_size' in config_updates or 'chunk_overlap' in config_updates:
            self.text_splitter = CharacterTextSplitter(
                chunk_size=self.config['chunk_size'],
                chunk_overlap=self.config['chunk_overlap']
            )
            # Mark that we need to re-process documents
            if self.uploaded_documents:
                # We should re-process, but for now just note it
                pass
        
        return {
            'success': True,
            'message': 'Configuration updated successfully',
            'updated_fields': list(config_updates.keys()),
            'current_config': self.config
        }
    
    @debug_track("Executing RAG Console Command")
    async def execute_console_command(self, command: str, api_key: str) -> dict:
        """Execute console commands for testing and debugging RAG system"""
        try:
            # Initialize components if needed
            await self._initialize_components(api_key)
            
            # Available commands and their implementations
            commands = {
                # Vector database commands
                'vector_db.get_stats': self._console_get_vector_stats,
                'vector_db.list_vectors': self._console_list_vectors,
                
                # Handler commands
                'handler.get_documents_info': self._console_get_documents_info,
                'handler.get_config': self._console_get_config,
                'handler.get_status': self._console_get_status,
                
                # Utility commands
                'help': self._console_help,
                'clear': self._console_clear,
                'status': self._console_get_status
            }
            
            # Parse command
            command = command.strip()
            
            if not command:
                return {"success": False, "output": "Error: Empty command", "command": command}
            
            # Handle parameterized commands (e.g., vector_db.search_by_text("query", k=3))
            if command.startswith('vector_db.search_by_text('):
                return await self._handle_search_command(command)
            elif command.startswith('handler.set_adjacent_chunks('):
                return await self._handle_set_adjacent_chunks(command)
            elif command.startswith('handler.compare_scoring('):
                return await self._handle_compare_scoring(command)
            elif command.startswith('handler.compare_all_methods('):
                return await self._handle_compare_all_methods(command)
            
            # Handle simple commands
            if command in commands:
                result = await commands[command]()
                return {"success": True, "output": result, "command": command}
            
            # Handle special cases
            if command.startswith('vector_db.'):
                return {"success": False, "output": f"Unknown vector_db command: {command}", "command": command}
            elif command.startswith('handler.'):
                return {"success": False, "output": f"Unknown handler command: {command}", "command": command}
            else:
                return {"success": False, "output": f"Unknown command: {command}. Type 'help' for available commands.", "command": command}
                
        except Exception as e:
            return {"success": False, "output": f"Error executing command: {str(e)}", "command": command}
    
    async def _handle_search_command(self, command: str) -> dict:
        """Handle vector_db.search_by_text() commands with parameters"""
        try:
            # Extract parameters from command like: vector_db.search_by_text("query", k=3)
            import re
            match = re.match(r'vector_db\.search_by_text\(\s*["\']([^"\']+)["\']\s*(?:,\s*k\s*=\s*(\d+))?\s*\)', command)
            
            if not match:
                return {"success": False, "output": "Invalid search command format. Use: vector_db.search_by_text(\"query\", k=5)", "command": command}
            
            query = match.group(1)
            k = int(match.group(2)) if match.group(2) else 5
            
            if not self.vector_db or len(getattr(self.vector_db, 'vectors', [])) == 0:
                return {"success": True, "output": "No vector database available. Upload documents first.", "command": command}
            
            # Use the main search method to get results with adjacent chunks if configured
            results = await self.search_documents(query, 'console-command', k=k)
            
            if not results:
                output = f"No results found for query: '{query}'"
            else:
                # Check if we have enhanced results with chunk info
                if isinstance(results, list) and results and 'chunk_info' in results[0]:
                    output_lines = [f"Found {len(results)} result(s) for query: '{query}'", ""]
                    for result in results:
                        chunk_info = result.get('chunk_info', {})
                        if 'center_chunk' in chunk_info:
                            # Enhanced results with adjacent chunks
                            center = chunk_info['center_chunk']
                            
                            # Ranking information
                            rank_info = f"Rank: {result['rank']}"
                            if 'original_rank' in result:
                                rank_change = result.get('rank_change', 0)
                                if rank_change > 0:
                                    rank_info += f" (was #{result['original_rank']}, ↑{rank_change})"
                                elif rank_change < 0:
                                    rank_info += f" (was #{result['original_rank']}, ↓{abs(rank_change)})"
                                else:
                                    rank_info += f" (unchanged from #{result['original_rank']})"
                            
                            output_lines.append(f"{result['rank']}. {rank_info}")
                            
                            # Score comparison
                            output_lines.append(f"   Original Score: {result.get('original_score', 0):.4f}")
                            output_lines.append(f"   Adjacent Score: {result['similarity_score']:.4f}")
                            
                            score_change = result.get('score_change', 0)
                            score_change_pct = result.get('score_change_pct', 0)
                            if score_change > 0:
                                output_lines.append(f"   Change: +{score_change:.4f} (+{score_change_pct:.1f}%)")
                            else:
                                output_lines.append(f"   Change: {score_change:.4f} ({score_change_pct:.1f}%)")
                            
                            # Document and chunk info
                            output_lines.append(f"   Document: {center['doc_name']}")
                            output_lines.append(f"   Chunk: {center['local_idx'] + 1} of {center['total_in_doc']} (with {chunk_info['adjacent_setting']} adjacent = {chunk_info['total_chunks_used']} total)")
                            
                            # Full text (no truncation)
                            output_lines.append(f"   Text: {result['text']}")
                        else:
                            # Standard results
                            output_lines.append(f"{result['rank']}. Score: {result['similarity_score']:.4f}")
                            if 'doc_name' in chunk_info:
                                output_lines.append(f"   Document: {chunk_info['doc_name']}")
                                output_lines.append(f"   Chunk: {chunk_info['local_idx'] + 1} of {chunk_info['total_in_doc']}")
                            output_lines.append(f"   Text: {result['text']}")
                        output_lines.append("")
                    
                    # Add summary if we have adjacent chunks
                    if any('center_chunk' in r.get('chunk_info', {}) for r in results):
                        output_lines.append("\n--- Adjacent Chunks Impact Summary ---")
                        
                        # Count ranking changes
                        improved = sum(1 for r in results if r.get('rank_change', 0) > 0)
                        worsened = sum(1 for r in results if r.get('rank_change', 0) < 0)
                        unchanged = sum(1 for r in results if r.get('rank_change', 0) == 0)
                        
                        output_lines.append(f"Ranking changes: ↑{improved} improved, ↓{worsened} dropped, ={unchanged} unchanged")
                        
                        # Average score change
                        avg_score_change = sum(r.get('score_change', 0) for r in results) / len(results)
                        avg_pct_change = sum(r.get('score_change_pct', 0) for r in results) / len(results)
                        output_lines.append(f"Average score change: {avg_score_change:+.4f} ({avg_pct_change:+.1f}%)")
                        
                        # Configuration info
                        output_lines.append(f"Configuration: adjacent_chunks={self.config.get('adjacent_chunks', 1)}, strategy={self.config.get('scoring_strategy', 'reembed')}")
                    
                    output = "\n".join(output_lines)
                else:
                    # Legacy format
                    output_lines = [f"Found {len(results)} result(s) for query: '{query}'", ""]
                    for i, (text, score) in enumerate(results, 1):
                        output_lines.append(f"{i}. Score: {score:.4f}")
                        output_lines.append(f"   Text: {text}")
                        output_lines.append("")
                    output = "\n".join(output_lines)
            
            return {"success": True, "output": output, "command": command}
            
        except Exception as e:
            return {"success": False, "output": f"Error in search command: {str(e)}", "command": command}
    
    async def _console_get_vector_stats(self) -> str:
        """Get vector database statistics"""
        if not self.vector_db:
            return "Vector database not initialized"
        
        stats = [
            f"Vector Database Status:",
            f"- Total vectors: {len(self.vector_db.vectors) if hasattr(self.vector_db, 'vectors') else 'Unknown'}",
            f"- Embedding model: {self.config['embedding_model']}",
            f"- Documents processed: {len(self.uploaded_documents)}",
            f"- Vector database ready: {self.vector_db is not None and len(getattr(self.vector_db, 'vectors', [])) > 0}"
        ]
        return "\n".join(stats)
    
    async def _console_list_vectors(self) -> str:
        """List first few vectors with previews"""
        if not self.vector_db or not hasattr(self.vector_db, 'vectors'):
            return "No vectors available"
        
        vectors = self.vector_db.vectors
        if not vectors:
            return "Vector database is empty"
        
        lines = [f"Vector Database Contents ({len(vectors)} total vectors):", ""]
        vector_list = list(vectors)[:5]  # Convert to list and take first 5
        for i, vector_data in enumerate(vector_list):  # Show first 5
            try:
                # Try to extract text content from vector data
                text_content = str(vector_data)[:100] + "..." if len(str(vector_data)) > 100 else str(vector_data)
                lines.append(f"{i+1}. {text_content}")
            except Exception:
                lines.append(f"{i+1}. [Vector data: {type(vector_data)}]")
        
        if len(vectors) > 5:
            lines.append(f"... and {len(vectors) - 5} more vectors")
        
        return "\n".join(lines)
    
    async def _console_get_documents_info(self) -> str:
        """Get information about uploaded documents"""
        if not self.uploaded_documents:
            return "No documents uploaded"
        
        lines = [f"Uploaded Documents ({len(self.uploaded_documents)} total):", ""]
        for i, doc in enumerate(self.uploaded_documents, 1):
            lines.append(f"{i}. {doc['name']} ({doc['type']})")
            lines.append(f"   Size: {doc['size']} chars, Words: {doc.get('word_count', 'N/A')}")
            lines.append(f"   Uploaded: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(doc['upload_time']))}")
            lines.append("")
        
        return "\n".join(lines)
    
    async def _console_get_config(self) -> str:
        """Get current RAG configuration"""
        lines = ["Current RAG Configuration:", ""]
        for key, value in self.config.items():
            lines.append(f"- {key}: {value}")
        return "\n".join(lines)
    
    async def _console_get_status(self) -> str:
        """Get overall RAG system status"""
        lines = [
            "RAG System Status:",
            f"- Documents uploaded: {len(self.uploaded_documents)}",
            f"- Text splitter ready: {self.text_splitter is not None}",
            f"- Embedding model ready: {self.embedding_model is not None}",
            f"- Vector database ready: {self.vector_db is not None}",
            f"- Chat model ready: {self.chat_model is not None}",
            ""
        ]
        
        if self.vector_db and hasattr(self.vector_db, 'vectors'):
            lines.append(f"- Total vectors in database: {len(self.vector_db.vectors)}")
        
        if self.uploaded_documents:
            total_size = sum(doc['size'] for doc in self.uploaded_documents)
            lines.append(f"- Total document size: {total_size} characters")
        
        return "\n".join(lines)
    
    async def _console_help(self) -> str:
        """Show available console commands"""
        return """RAG Console - Available Commands:

Vector Database:
  vector_db.search_by_text("query", k=5)  - Search for similar text chunks
  vector_db.get_stats()                   - Show vector database statistics  
  vector_db.list_vectors()                - List first few vectors with previews

Handler Info:
  handler.get_documents_info()            - Show uploaded document details
  handler.get_config()                    - Display current RAG configuration
  handler.get_status()                    - Show overall system status

Adjacent Chunks Testing:
  handler.set_adjacent_chunks(n)          - Set adjacent chunks (0-5)
  handler.compare_scoring("query", k=5)   - Compare normal vs adjacent scoring
  handler.compare_all_methods("query", k=5, chars=200) - Compare all 3 chunking strategies

Utilities:
  help                                    - Show this help message
  status                                  - Quick status check
  clear                                   - Clear console (frontend only)

Examples:
  vector_db.search_by_text("machine learning", k=3)
  handler.set_adjacent_chunks(2)
  handler.compare_scoring("neural networks", k=5)
  handler.compare_all_methods("AI research", k=5, chars=150)
  handler.get_config()"""
    
    async def _console_clear(self) -> str:
        """Clear console (handled by frontend)"""
        return "Console cleared"
    
    async def _handle_set_adjacent_chunks(self, command: str) -> dict:
        """Handle handler.set_adjacent_chunks(n) commands"""
        try:
            import re
            match = re.match(r'handler\.set_adjacent_chunks\(\s*(\d+)\s*\)', command)
            
            if not match:
                return {"success": False, "output": "Invalid format. Use: handler.set_adjacent_chunks(n) where n is 0-5", "command": command}
            
            n = int(match.group(1))
            
            if n < 0 or n > 5:
                return {"success": False, "output": "Adjacent chunks must be between 0 and 5", "command": command}
            
            old_value = self.config.get('adjacent_chunks', 1)
            self.config['adjacent_chunks'] = n
            
            output = f"Adjacent chunks setting changed from {old_value} to {n}\n"
            if n == 0:
                output += "Adjacent chunks disabled - using single chunks only"
            else:
                total_chunks = 2 * n + 1
                output += f"Now using {total_chunks} total chunks ({n} on each side of the match)"
            
            return {"success": True, "output": output, "command": command}
            
        except Exception as e:
            return {"success": False, "output": f"Error setting adjacent chunks: {str(e)}", "command": command}
    
    async def _handle_compare_scoring(self, command: str) -> dict:
        """Handle handler.compare_scoring("query", k=5) commands"""
        try:
            import re
            match = re.match(r'handler\.compare_scoring\(\s*["\']([^"\']+)["\']\s*(?:,\s*k\s*=\s*(\d+))?\s*\)', command)
            
            if not match:
                return {"success": False, "output": "Invalid format. Use: handler.compare_scoring(\"query\", k=5)", "command": command}
            
            query = match.group(1)
            k = int(match.group(2)) if match.group(2) else 5
            
            if not self.vector_db or len(getattr(self.vector_db, 'vectors', [])) == 0:
                return {"success": True, "output": "No vector database available. Upload documents first.", "command": command}
            
            current_adjacent = self.config.get('adjacent_chunks', 1)
            
            # First search with adjacent_chunks = 0 (normal)
            self.config['adjacent_chunks'] = 0
            normal_results = await self.search_documents(query, 'console-command', k=k)
            
            # Then search with the configured adjacent chunks
            self.config['adjacent_chunks'] = current_adjacent
            adjacent_results = await self.search_documents(query, 'console-command', k=k)
            
            # Format comparison output
            output_lines = [f"Scoring Comparison for query: '{query}' (k={k})", ""]
            output_lines.append(f"Normal scoring (adjacent_chunks=0) vs Adjacent scoring (adjacent_chunks={current_adjacent})")
            output_lines.append("=" * 80)
            
            for i in range(min(len(normal_results), len(adjacent_results))):
                normal = normal_results[i]
                adjacent = adjacent_results[i] if i < len(adjacent_results) else None
                
                output_lines.append(f"\nPosition {i+1}:")
                output_lines.append(f"  Normal:   Score {normal['similarity_score']:.4f} - {normal['text']}")
                
                if adjacent:
                    rank_change = ""
                    if 'original_rank' in adjacent:
                        orig_rank = adjacent['original_rank']
                        if orig_rank != i + 1:
                            if orig_rank > i + 1:
                                rank_change = f" (↑ from #{orig_rank})"
                            else:
                                rank_change = f" (↓ from #{orig_rank})"
                    
                    output_lines.append(f"  Adjacent: Score {adjacent['similarity_score']:.4f}{rank_change} - {adjacent['text']}")
                    
                    if 'score_change' in adjacent:
                        change = adjacent['score_change']
                        pct_change = adjacent.get('score_change_pct', 0)
                        output_lines.append(f"  Change:   {change:+.4f} ({pct_change:+.1f}%)")
                else:
                    output_lines.append("  Adjacent: [No result at this position]")
            
            return {"success": True, "output": "\n".join(output_lines), "command": command}
            
        except Exception as e:
            return {"success": False, "output": f"Error in scoring comparison: {str(e)}", "command": command}
    
    async def _handle_compare_all_methods(self, command: str) -> dict:
        """Handle handler.compare_all_methods("query", k=5, chars=None) commands"""
        try:
            import re
            # Updated regex to capture k and chars parameters
            match = re.match(r'handler\.compare_all_methods\(\s*["\']([^"\']+)["\']\s*(?:,\s*k\s*=\s*(\d+))?(?:,\s*chars\s*=\s*(\d+))?\s*\)', command)
            
            if not match:
                return {"success": False, "output": "Invalid format. Use: handler.compare_all_methods(\"query\", k=5, chars=200)", "command": command}
            
            query = match.group(1)
            k = int(match.group(2)) if match.group(2) else 5
            chars = int(match.group(3)) if match.group(3) else None  # None means show full text
            
            if not self.vector_db or len(getattr(self.vector_db, 'vectors', [])) == 0:
                return {"success": True, "output": "No vector database available. Upload documents first.", "command": command}
            
            current_adjacent = self.config.get('adjacent_chunks', 1)
            
            def format_text(text: str, char_limit: Optional[int] = None) -> str:
                """Format text with optional character limit"""
                if char_limit is None:
                    return text
                if len(text) <= char_limit:
                    return text
                return text[:char_limit] + "..."
            
            # Method 1: Normal small chunks
            self.config['adjacent_chunks'] = 0
            normal_results = await self.search_documents(query, 'console-command', k=k)
            
            # Method 2: Small chunks + adjacent expansion
            self.config['adjacent_chunks'] = current_adjacent
            adjacent_results = await self.search_documents(query, 'console-command', k=k)
            
            # Method 3: Large chunks from the start
            large_results = []
            if self.large_chunk_vector_db and len(getattr(self.large_chunk_vector_db, 'vectors', [])) > 0:
                large_search_results = self.large_chunk_vector_db.search_by_text(query, k=k)
                for i, (text, score) in enumerate(large_search_results):
                    # Find metadata for large chunk
                    chunk_idx = self._find_large_chunk_index_by_text(text)
                    metadata = self.large_chunk_metadata.get(chunk_idx, {})
                    
                    large_results.append({
                        'rank': i + 1,
                        'text': text,
                        'similarity_score': float(score),
                        'full_text': text,
                        'chunk_info': {
                            'doc_name': metadata.get('doc_name', 'Unknown'),
                            'local_idx': metadata.get('local_chunk_idx', -1),
                            'total_in_doc': metadata.get('total_chunks_in_doc', -1),
                            'chunk_size': metadata.get('chunk_size', len(text))
                        }
                    })
            
            # Format comprehensive comparison output
            output_lines = [f"📊 COMPREHENSIVE CHUNKING STRATEGY COMPARISON", ""]
            output_lines.append(f"Query: '{query}' (k={k})")
            output_lines.append(f"Small chunks: {self.config['chunk_size']} chars")
            output_lines.append(f"Adjacent expansion: {current_adjacent} chunks each side")
            output_lines.append(f"Large chunks: {self.config['comparison_chunk_size']} chars")
            output_lines.append("=" * 100)
            
            for i in range(k):
                output_lines.append(f"\n🏆 RANK #{i+1} COMPARISON")
                output_lines.append("-" * 50)
                
                # Method 1: Small chunks
                if i < len(normal_results):
                    normal = normal_results[i]
                    output_lines.append(f"📄 SMALL CHUNKS: Score {normal['similarity_score']:.4f}")
                    output_lines.append(f"   {format_text(normal['text'], chars)}")
                else:
                    output_lines.append(f"📄 SMALL CHUNKS: [No result at rank {i+1}]")
                
                output_lines.append("")
                
                # Method 2: Adjacent expansion
                if i < len(adjacent_results):
                    adjacent = adjacent_results[i]
                    chunk_info = adjacent.get('chunk_info', {}).get('center_chunk', {})
                    total_chunks_used = adjacent.get('chunk_info', {}).get('total_chunks_used', 1)
                    
                    # Show ranking change for adjacent method
                    rank_change_info = ""
                    if 'original_rank' in adjacent:
                        orig_rank = adjacent['original_rank']
                        if orig_rank != i + 1:
                            if orig_rank > i + 1:
                                rank_change_info = f" (was #{orig_rank}, ↑{orig_rank - (i + 1)})"
                            else:
                                rank_change_info = f" (was #{orig_rank}, ↓{(i + 1) - orig_rank})"
                    
                    output_lines.append(f"🔗 ADJACENT EXPANSION: Score {adjacent['similarity_score']:.4f}{rank_change_info}")
                    output_lines.append(f"   {format_text(adjacent['text'], chars)}")
                    output_lines.append(f"   (Using {total_chunks_used} chunks: {current_adjacent} adjacent on each side)")
                else:
                    output_lines.append(f"🔗 ADJACENT EXPANSION: [No result at rank {i+1}]")
                
                output_lines.append("")
                
                # Method 3: Large chunks
                if i < len(large_results):
                    large = large_results[i]
                    chunk_info = large.get('chunk_info', {})
                    
                    output_lines.append(f"📚 LARGE CHUNKS: Score {large['similarity_score']:.4f}")
                    output_lines.append(f"   {format_text(large['text'], chars)}")
                    output_lines.append(f"   (Single large chunk: {chunk_info.get('chunk_size', 'Unknown')} chars)")
                else:
                    output_lines.append(f"📚 LARGE CHUNKS: [No result at rank {i+1}]")
            
            # Overall summary
            output_lines.append(f"\n📈 OVERALL ANALYSIS")
            output_lines.append("-" * 40)
            
            # Calculate average scores
            normal_avg = sum(r['similarity_score'] for r in normal_results[:k]) / min(k, len(normal_results)) if normal_results else 0
            adjacent_avg = sum(r['similarity_score'] for r in adjacent_results[:k]) / min(k, len(adjacent_results)) if adjacent_results else 0
            large_avg = sum(r['similarity_score'] for r in large_results[:k]) / min(k, len(large_results)) if large_results else 0
            
            output_lines.append(f"Average Scores (top {k}):")
            output_lines.append(f"  Small Chunks:      {normal_avg:.4f}")
            output_lines.append(f"  Adjacent Expanded: {adjacent_avg:.4f} ({adjacent_avg - normal_avg:+.4f})")
            output_lines.append(f"  Large Chunks:      {large_avg:.4f} ({large_avg - normal_avg:+.4f})")
            
            # Determine overall winner
            best_method = max([
                ('Small', normal_avg),
                ('Adjacent', adjacent_avg),
                ('Large', large_avg)
            ], key=lambda x: x[1])
            
            output_lines.append(f"\n🏆 BEST OVERALL METHOD: {best_method[0]} (avg score: {best_method[1]:.4f})")
            
            return {"success": True, "output": "\n".join(output_lines), "command": command}
            
        except Exception as e:
            return {"success": False, "output": f"Error in comprehensive comparison: {str(e)}", "command": command}
    
    def _find_large_chunk_index_by_text(self, text: str) -> int:
        """Find large chunk index by matching text content"""
        for idx, metadata in self.large_chunk_metadata.items():
            if metadata['chunk_text'] == text:
                return idx
        return -1